import gspread
import time
from oauth2client.service_account import ServiceAccountCredentials
from pprint import pprint
import docx
from docx.shared import Pt


document = docx.Document()
style = document.styles['Normal']
font = style.font
font.name = 'TimesNewRoman'
font.size = Pt(12)

p = document.add_paragraph()
r = p.add_run('Список представляемых к публикации докладов\n')
r.bold = True
r.italic = True
p.alignment = 1


document.add_paragraph('\tКафедра № 43 компьютерных технологий и программной инженерии')
document.add_paragraph('\tПоляк Марк Дмитриевич')
document.add_paragraph('\te-mail: m.polyak@guap.ru')
document.add_paragraph('\tтел.: +7-999-XXX-XXXX\n')



ex = []
delay1 = 60 #should be 60sec
delay2 = 0.1
q=0


try:

    try:
        time.sleep(delay1) #задержка перед вызовом, для предотвращения исключений
    except RuntimeError:
        ex.append(RuntimeError)


    scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
             "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]

    creds = ServiceAccountCredentials.from_json_keyfile_name("creds.json", scope)

    client = gspread.authorize(creds)

    sheet = client.open("МСНК-2019").sheet1  # Open the spreadhseet

    data = sheet.get_all_records()
    q = len(data)
except gspread.exceptions.APIError:
    pprint("Квота Google Sheets API превышена.")
res = []
j=0
k=1



for i in range (2, q+2):
    try:
        j=j+1
        if(j>=90):
            j=0
            try:
                time.sleep(delay1)  # задержка перед вызовом, для предотвращения исключений
            except RuntimeError:
                ex.append(RuntimeError)

        try:
            time.sleep(delay2)  # задержка перед вызовом, для предотвращения исключений
        except RuntimeError:
            ex.append(RuntimeError)
        row = sheet.row_values(i)

        try:
            if((row[10] == '1' or row[10] == 'Да') and (row[11] == '1' or row[11] == 'Да')):
                s = str(k) + '.\t' + row[0] + " " + row[1][0] + ". " + row[2][0] + ". " + row[12]
                document.add_paragraph('\t'+s)
                k=k+1
        except IndexError:
            ex.append(IndexError) #Таблицу заполнить нужно!"

    except gspread.exceptions.APIError:
        pprint("Квота2 Google Sheets API превышена.")
        q=q-1
        try:
            time.sleep(delay1)  # задержка перед вызовом, для предотвращения исключений
        except RuntimeError:
            ex.append(RuntimeError)


pprint(ex)

try:
    document.save('Список представляемых к публикации докладов.docx')
except PermissionError:
    pprint("Файл закрыть нужно перед перезаписью. PermissionError: Permission denied: 'Список представляемых к публикации докладов.docx'")

