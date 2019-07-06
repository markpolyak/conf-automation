from datetime import datetime
from docx import Document
from docx.enum.text import *
from docx.shared import *
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import sys

def main():
    rowData = getData()
    createReport(getMeetingMembers(rowData))

def createReport(data):
    """
    Функция формирует отчет о заседаниях МСНК и сохраняет его в файл [fileName].docx.
    Входными данными являются результаты заседаний, находящиеся в словаре «data».
    Ключ словаря -  номер заседания.
    Значение словаря - список списков со значениями для таблицы 
    """
    #==================== Настройки документа ===================
    print('Генерация отчёта.')
    document = Document()
    fileName = 'report.docx'
    n = '72-й' #номер конференции
    #Стиль документа
    style = document.styles['Normal']
    font = style.font
    font.name = 'TimesNewRoman'
    font.size = Pt(10)
    #Настройка межстрочного интервала
    style.paragraph_format.space_before = Pt(0) 
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #==================== Формирование документа ===================
    #шапка
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Отчет о проведении ' + n + ' МСНК ГУАП')
    r.bold = True
    p = document.add_paragraph()    
    r = p.add_run('Секция 43. Компьютерных технологий и программной инженерии')
    r.bold = True
    p = document.add_paragraph()
    #вставка таблиц
    for i in data:
        addMeeting(i,document,data)   
    #Подпись
    p = document.add_paragraph()
    r = p.add_run('Научный руководитель секции                             _________________ / Охтилев М.Ю.')
    #================ Сохранение документа ======================
    try :
        document.save(fileName)
        print('Отчёт сохранён в файле: ' + fileName)
    except Exception as err:
        print('Ошибка сохранения файла: ' + str(err))
        fileName = datetime.now().strftime("%Y-%m-%d_%H-%M-%S_") + fileName
        document.save(fileName)
        print('Отчёт сохранён в файле: ' + fileName)

def addMeeting(l,document,data):
    """
    Функция добавляет заседание в отчёт
    l – номер заседания 
    document – отчёт 
    data  - список докладов 
    """
    #настройка таблиц
    tbFont = 'Arial' 
    tbStyle = 'Table Grid'
    #ширина колонок таблиц
    col1width = Cm(0.77)
    col2width = Cm(8.81)
    col3width = Cm(2.35)
    col4width = Cm(4.34)
    #шапка заседания
    print("Вставляем заседание " + str(l))
    p = document.add_paragraph()
    r = p.add_run('Заседание ' + str(l))
    r.bold = True
    p = document.add_paragraph()
    r = p.add_run('MeetingDateTime' + '\t\t\t\t' + 'ул. Б. Морская, д. 67, ауд. ' + 'MeetingRoom')
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p = document.add_paragraph()
    r = p.add_run('Научный руководитель секции – д-р техн. наук, проф. М.Ю. Охтилев')
    p = document.add_paragraph()
    r = p.add_run('Секретарь – ' + 'secretary')
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p = document.add_paragraph()
    p = document.add_paragraph()
    r = p.add_run('\tСписок докладов')
    p = document.add_paragraph()

    createTB_s = datetime.now()

    #таблица с докладами
    #шапка таблицы
    table1 = document.add_table(rows=1, cols=4)
    table1.style = tbStyle
    cells = table1.rows[0].cells
    cells[0].paragraphs[0].add_run('№ п/п')
    cells[1].paragraphs[0].add_run('Фамилия и инициалы докладчика, название доклада')
    cells[2].paragraphs[0].add_run('Статус (магистр / студент)')
    cells[3].paragraphs[0].add_run('Решение')
    for cell in table1.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    #заполнение таблицы
    fill = datetime.now()
    for i in data.get(l):
        row_cells = table1.add_row().cells
        row_cells[0].text = i[0]
        row_cells[1].text = i[1]
        row_cells[2].text = i[2]
        row_cells[3].text = i[3]
    print("Заполнение таблицы " + str(datetime.now() - fill))

    wit = datetime.now()
    col = table1.columns
    for cel in col[0].cells:
        cel.width = col1width
    for cel in col[1].cells:
        cel.width = col2width
    for cel in col[2].cells:
        cel.width = col3width
    for cel in col[3].cells:
        cel.width = col4width
    print("Установка размера " + str(datetime.now() - wit))

    ft = datetime.now()
    for row in table1.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.name = tbFont
    print("Установка шрифта " + str(datetime.now() - ft))
    
    p = document.add_paragraph()
    p = document.add_paragraph()

    print("Создание таблицы " + str(datetime.now() - createTB_s))
    
def getData():
    """
    Функция получает данные из Гугл таблицы с результатами проведения конференции.
    """
    print("Получение данных.")
    s = datetime.now()
    scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
             "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_name("key2.json",scope) #key2.json файл для авторизации
    client = gspread.authorize(creds)
    try:
        sheet = client.open("МСНК-2019").sheet1 #таблица с данными
        tmp = sheet.get_all_values()
        print("Данные получены за: " + str(datetime.now() - s))
    except Exception as err:
        print("Ошибка получения данных :" + err)
        sys.exit(-1)
    return tmp

def getMeetingMembers(rowData):
    """
    Функция получает на вход содержимое гугл таблицы с докладами.
    Возвращает словарь со списками выступлений на заседаниях (ключ номер заседания) 
    """
    print("Подготовка данных.")
    s = datetime.now()
    data = {}
    tb = []
    n = 0
    k = 1
    print("Сортируем доклады по дате.")
    tmp = sorted(rowData,key=lambda row:row[8]) #сортируем доклады по дате
    for j in range(0,len(tmp)):
            if tmp[j][8] != "":
                curDate = tmp[j][8]
                break
    print("Считаем заседания и формируем списки участников.")
    for i in range(j,len(tmp)):
         if tmp[i][8] != "" and tmp[i][0] != "Фамилия":
            if tmp[i][8] != curDate:
                curDate = tmp[i][8]
                data.update({k:tb})
                k+=1
                n = 0
                tb = []
            if tmp[i][9] == "1" or tmp[i][9] == "да":
                n+=1
                tb.append([str(n),tmp[i][0] + " " + tmp[i][1] + " " + tmp[i][2] + "." + tmp[i][4],tmp[i][3],"опубликовать доклад в сборнике СНК"])
            else :
                n+=1
                tb.append([str(n),tmp[i][0] + " " + tmp[i][1] + " " + tmp[i][2] + "." + tmp[i][4],tmp[i][3],"доклад плохо подготовлен"])
    data.update({k:tb})
    print("Данные подготовлены за: " + str(datetime.now() - s))
    return data      

if __name__ == "__main__":
    main()